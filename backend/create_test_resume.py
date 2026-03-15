
from docx import Document
import io

def create_resume_docx():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('john.doe@email.com | 123-456-7890 | github.com/johndoe | linkedin.com/in/johndoe')
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Computer Science student with experience in Python and React.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.Tech in Computer Science Engineering, 2020-2024')
    doc.add_paragraph('CGPA: 8.5')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, JavaScript, React, SQL, Git, Data Structures, Algorithms')
    
    doc.add_heading('Projects', level=1)
    doc.add_paragraph('Project 1: Personal Portfolio - Built using React and Node.js')
    doc.add_paragraph('Project 2: E-commerce Site - Python/Django backend')
    
    doc.save('test_resume.docx')

if __name__ == "__main__":
    create_resume_docx()
