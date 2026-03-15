# -------------------------------------------------------
import re
import io
import logging
from typing import Any, Optional

logger = logging.getLogger(__name__)


def extract_text_with_pymupdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # type: ignore
        text_parts = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_text = page.get_text("text")
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
        doc.close()
        return "\n".join(text_parts)
    except Exception:
        return ""


def extract_text_with_pdfplumber(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber."""
    try:
        import pdfplumber  # type: ignore
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception:
        return ""


def extract_text_with_ocr(file_bytes: bytes) -> str:
    """
    OCR fallback for scanned/image-based PDFs.
    Uses PyMuPDF to render pages as images, then RapidOCR to read them.
    No external binaries needed — everything is pure Python.
    """
    try:
        import fitz  # type: ignore
        from rapidocr_onnxruntime import RapidOCR  # type: ignore
        from PIL import Image  # type: ignore
        import numpy as np  # type: ignore

        ocr_engine = RapidOCR()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        all_text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            # Render page as image at 300 DPI for good OCR quality
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image then to numpy array for RapidOCR
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img_array = np.array(img)

            # Run OCR
            result, _ = ocr_engine(img_array)
            if result:
                page_texts = [line[1] for line in result if line[1]]
                if page_texts:
                    all_text_parts.append("\n".join(page_texts))

        doc.close()
        return "\n\n".join(all_text_parts)

    except ImportError as e:
        print(f"OCR dependencies missing: {e}")
        return ""
    except Exception as e:
        print(f"OCR extraction failed: {e}")
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using multiple methods:
    1. PyMuPDF (fitz) — fast, handles most PDFs
    2. pdfplumber — good for table-heavy PDFs
    3. RapidOCR — for scanned/image-based PDFs (OCR)
    """
    # Method 1: PyMuPDF
    text = extract_text_with_pymupdf(file_bytes)
    if text and len(text.split()) >= 10:
        print(f"[Parser] PyMuPDF extracted {len(text.split())} words")
        return text

    # Method 2: pdfplumber
    text = extract_text_with_pdfplumber(file_bytes)
    if text and len(text.split()) >= 10:
        print(f"[Parser] pdfplumber extracted {len(text.split())} words")
        return text

    # Method 3: OCR fallback for scanned PDFs
    print("[Parser] Scanned/Image-based PDF detected. Initializing OCR engine...")
    text = extract_text_with_ocr(file_bytes)
    if text and len(text.split()) >= 3:
        print(f"[Parser] OCR successfully extracted {len(text.split())} words")
        return text

    return text or ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file including paragraphs, tables, headers, and text boxes."""
    import docx  # type: ignore
    from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
    
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        print(f"[Parser] DOCX loaded. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
        text_parts = []

        # 1. Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # 2. Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        # 3. Extract from Text Boxes (Shapes) - often missed by doc.paragraphs
        # This is a bit advanced as text boxes are deeply nested in XML
        try:
            for rel in doc.part.rels.values():
                if rel.reltype == RT.HEADER or rel.reltype == RT.FOOTER:
                    header_footer_doc = docx.Document(io.BytesIO(rel.target_part.blob))
                    for p in header_footer_doc.paragraphs:
                        if p.text.strip():
                            text_parts.append(p.text.strip())
        except:
            pass
            
        # Recursive XML search for text (fallback for text boxes / complex shapes)
        def deep_xml_extract(element) -> str:
            text = ""
            if type(element.tag) is str and element.tag.endswith('t'): # <w:t> is the standard text tag
                text += str(element.text) if element.text else ""
            for child in element:
                text += str(deep_xml_extract(child))
            # Add space between major elements if needed
            if type(element.tag) is str and (element.tag.endswith('p') or element.tag.endswith('tr')):
                text += "\n"
            return text
            
        if not text_parts or len(" ".join(text_parts).split()) < 20:
            print("[Parser] Standard extraction incomplete, performing deep XML text search...")
            xml_text = deep_xml_extract(doc.element)
            if xml_text.strip():
                text_parts.append(xml_text.strip())

        print(f"[Parser] DOCX extraction complete. Parts: {len(text_parts)}")
        final_text = "\n".join(text_parts)
        return final_text
    except Exception as e:
        print(f"[Parser] DOCX main error: {e}")
        import traceback
        traceback.print_exc()
        return ""


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s@.\-+/|:,()%]', ' ', text)
    return text.strip()


def split_into_chunks(text: str, chunk_size: int = 200, overlap: int = 50) -> list[str]:
    """Split resume text into overlapping chunks for RAG embedding."""
    words = text.split()
    chunks: list[str] = []
    i = 0
    word_count = len(words)
    while i < word_count:
        # FIXED PROBLEM: Explicit slicing for picky linter
        end_idx = min(i + chunk_size, word_count)
        chunk_words = [words[j] for j in range(i, end_idx)]
        if chunk_words:
            chunks.append(" ".join(chunk_words))
        i += chunk_size - overlap
        if i >= word_count:
            break
    return chunks


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """Detect file type and extract text using best available method."""
    filename_lower = filename.lower()
    raw_text = ""
    file_type = "unknown"

    # Deep Signature Check (rare but happens: PDF renamed to DOCX)
    is_pdf = file_bytes.startswith(b"%PDF-")

    if is_pdf or filename_lower.endswith(".pdf"):
        print(f"[Parser] Processing as PDF (Signature check: {is_pdf})")
        raw_text = extract_text_from_pdf(file_bytes)
        file_type = "pdf"
    elif filename_lower.endswith(".docx"):
        print(f"[Parser] Processing as DOCX")
        raw_text = extract_text_from_docx(file_bytes)
        file_type = "docx"
    
    # Final backup: If extraction failed but it's a small file, try the other way just in case
    if not raw_text.strip() and not is_pdf:
        print("[Parser] DOCX extraction failed, trying PDF fallback as last resort...")
        raw_text = extract_text_from_pdf(file_bytes)
        if raw_text: file_type = "pdf"

    if not raw_text.strip():
        raise ValueError(
            "Could not extract text from resume. The file may be corrupted, empty, or "
            "it might only contain images (like a photo of a resume). "
            "Please try saving your resume as a standard PDF and upload again."
        )

    clean = clean_text(raw_text)
    word_count = len(clean.split())

    if word_count < 20:
        raise ValueError(
            f"Resume too short ({word_count} words). Please upload a more complete resume."
        )

    # Chunks are generated from the cleaned text
    chunks = split_into_chunks(clean)

    return {
        "raw_text": raw_text,
        "cleaned_text": clean,
        "chunks": chunks,
        "word_count": word_count,
        "char_count": len(raw_text), # Use raw_text length for char_count
        "file_type": file_type
    }