# app/template_matcher.py
# -------------------------------------------------------
# Compares user's resume vs your college resume template
# Keyword-by-keyword + section-by-section match
# -------------------------------------------------------

import re
import docx
import io
from pathlib import Path


# -------------------------------------------------------
# Template loader (DOCX)
# -------------------------------------------------------

def load_template_from_docx(template_path: str) -> dict:
    """
    Load and parse your DOCX resume template.
    Extracts all keywords and sections.
    """
    doc = docx.Document(template_path)
    full_text = []
    sections_found = {}
    current_section = "general"

    # Section header detection patterns
    section_patterns = {
        "contact": r"(contact|personal|address|email|phone)",
        "summary": r"(summary|objective|profile|about)",
        "skills": r"(skill|technical|technologies|tools)",
        "experience": r"(experience|work|employment|internship)",
        "education": r"(education|academic|qualification|degree)",
        "projects": r"(project|work done|portfolio)",
        "certifications": r"(certification|certificate|course|training)",
        "achievements": r"(achievement|award|honor|activity)",
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect if this is a section header
        for section_name, pattern in section_patterns.items():
            if re.search(pattern, text.lower()) and len(text) < 50:
                current_section = section_name
                break

        if current_section not in sections_found:
            sections_found[current_section] = []
        sections_found[current_section].append(text)
        full_text.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())

    template_text = " ".join(full_text).lower()

    # Extract all meaningful keywords from template
    words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#.]*\b', template_text)
    # Filter stop words
    stop_words = {
        "the", "a", "an", "and", "or", "but", "in", "on", "at", "to",
        "for", "of", "with", "by", "from", "is", "was", "are", "be",
        "been", "have", "has", "had", "do", "does", "did", "will",
        "would", "could", "should", "may", "might", "shall", "not",
        "this", "that", "these", "those", "your", "my", "his", "her",
        "its", "our", "their", "what", "which", "who", "how", "when",
        "where", "why", "all", "any", "both", "each", "few", "more",
        "most", "other", "some", "such", "than", "too", "very", "just"
    }
    keywords = [w.lower() for w in words if w.lower() not in stop_words and len(w) > 2]
    unique_keywords = list(set(keywords))

    return {
        "template_text": template_text,
        "template_keywords": unique_keywords,
        "sections": sections_found,
        "total_keywords": len(unique_keywords)
    }


def load_template_from_bytes(template_bytes: bytes) -> dict:
    """Load template from bytes (for API upload)."""
    doc = docx.Document(io.BytesIO(template_bytes))
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())

    template_text = " ".join(full_text).lower()
    words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#.]*\b', template_text)
    stop_words = {
        "the", "a", "an", "and", "or", "but", "in", "on", "at", "to",
        "for", "of", "with", "by", "from", "is", "was", "are", "be",
        "been", "have", "has", "had", "do", "does", "did", "will",
        "would", "could", "should", "may", "might", "shall", "not",
        "this", "that", "these", "those", "your", "my", "his", "her",
        "its", "our", "their", "what", "which", "who", "how", "when",
        "where", "why", "all", "any", "both", "each", "few", "more",
        "most", "other", "some", "such", "than", "too", "very", "just"
    }
    keywords = [w.lower() for w in words if w.lower() not in stop_words and len(w) > 2]
    unique_keywords = list(set(keywords))

    return {
        "template_text": template_text,
        "template_keywords": unique_keywords,
        "total_keywords": len(unique_keywords)
    }


# -------------------------------------------------------
# Template vs Resume matching
# -------------------------------------------------------

def match_resume_to_template(resume_text: str, template_data: dict) -> dict:
    """
    Compare user resume against your template.
    Returns keyword-by-keyword match report.
    """
    resume_lower = resume_text.lower()
    template_keywords = template_data["template_keywords"]

    matched = []
    missing = []

    for keyword in template_keywords:
        pattern = r'\b' + re.escape(keyword) + r'\b'
        if re.search(pattern, resume_lower):
            matched.append(keyword)
        else:
            missing.append(keyword)

    total = len(template_keywords)
    match_count = len(matched)
    match_percentage = round((match_count / total) * 100, 1) if total > 0 else 0

    # Grade
    if match_percentage >= 80:
        grade = "Excellent Match"
        color = "green"
    elif match_percentage >= 60:
        grade = "Good Match"
        color = "blue"
    elif match_percentage >= 40:
        grade = "Partial Match"
        color = "yellow"
    else:
        grade = "Low Match"
        color = "red"

    # Prioritize important missing keywords (longer = more specific = more important)
    priority_missing = sorted(missing, key=len, reverse=True)[:20]

    return {
        "template_match_score": match_percentage,
        "grade": grade,
        "color": color,
        "matched_count": match_count,
        "total_template_keywords": total,
        "matched_keywords": matched[:30],       # Top 30 matched
        "missing_keywords": priority_missing,    # Top 20 most important missing
        "all_missing_count": len(missing),
        "recommendation": get_match_recommendation(match_percentage)
    }


def get_match_recommendation(score: float) -> str:
    if score >= 80:
        return "Your resume closely matches our college template. You're well-prepared!"
    elif score >= 60:
        return "Good template alignment. Add the missing keywords to strengthen your resume."
    elif score >= 40:
        return "Several key sections from the template are missing. Review the template carefully."
    else:
        return "Your resume needs significant restructuring. Use our college template as your base."


def count_projects(resume_text: str) -> dict:
    """Estimate number of projects mentioned in resume."""
    text_lower = resume_text.lower()

    # Count project section items
    project_indicators = [
        r'\bproject\s*\d+\b',           # "Project 1", "Project 2"
        r'\b\d+\.\s+[A-Z]',             # "1. ProjectName"
        r'github\.com/[\w-]+/[\w-]+',   # GitHub repo links
        r'\|\s*[A-Z][a-z]+.*?\|',       # "| React | Node.js |" style
    ]

    count = 0
    project_names = []

    for pattern in project_indicators:
        matches = re.findall(pattern, resume_text)
        count += len(matches)

    # Simple heuristic: count bullet points in project section
    lines = resume_text.split('\n')
    in_project_section = False
    for line in lines:
        line_lower = line.lower().strip()
        if 'project' in line_lower and len(line) < 30:
            in_project_section = True
        elif in_project_section and any(x in line_lower for x in ['education', 'skills', 'experience', 'certif']):
            in_project_section = False
        elif in_project_section and line.strip() and len(line.strip()) > 10:
            project_names.append(line.strip()[:60])

    estimated_count = max(count, len([p for p in project_names if len(p) > 15]) // 3)
    estimated_count = min(estimated_count, 15)  # Cap at 15

    return {
        "estimated_count": estimated_count,
        "status": get_project_status(estimated_count),
        "target": "5-10 projects recommended for placement"
    }


def get_project_status(count: int) -> str:
    if count >= 8:
        return "Excellent - Strong project portfolio"
    elif count >= 5:
        return "Good - Meeting placement target"
    elif count >= 3:
        return "Average - Add 2-3 more projects"
    elif count >= 1:
        return "Low - Need more projects (target: 5+)"
    else:
        return "Missing - No projects detected!"


def check_portfolio_links(resume_text: str) -> dict:
    """Check for GitHub, LinkedIn, Portfolio links."""
    text_lower = resume_text.lower()

    checks = {
        "github": {
            "found": bool(re.search(r'github\.com', text_lower)),
            "pattern": r'github\.com/[\w-]+',
            "importance": "High - Shows real code to recruiters"
        },
        "linkedin": {
            "found": bool(re.search(r'linkedin\.com', text_lower)),
            "pattern": r'linkedin\.com/in/[\w-]+',
            "importance": "High - Professional networking"
        },
        "portfolio": {
            "found": bool(re.search(r'(portfolio|website|personal site|vercel|netlify|github\.io)', text_lower)),
            "pattern": None,
            "importance": "Medium - Shows frontend/design skills"
        },
        "kaggle": {
            "found": bool(re.search(r'kaggle\.com', text_lower)),
            "pattern": r'kaggle\.com/[\w-]+',
            "importance": "High for AI/DS - Shows ML projects"
        },
        "leetcode": {
            "found": bool(re.search(r'leetcode\.com', text_lower)),
            "pattern": r'leetcode\.com/[\w-]+',
            "importance": "Medium - Shows DSA practice"
        }
    }

    present = [k for k, v in checks.items() if v["found"]]
    absent = [k for k, v in checks.items() if not v["found"]]

    portfolio_score = round((len(present) / len(checks)) * 100, 1)

    return {
        "portfolio_score": portfolio_score,
        "present": present,
        "absent": absent,
        "details": checks,
        "recommendation": get_portfolio_recommendation(present, absent)
    }


def get_portfolio_recommendation(present: list, absent: list) -> str:
    critical_missing = [x for x in ["github", "linkedin"] if x in absent]
    if critical_missing:
        return f"⚠️ Critical: Add {', '.join(critical_missing)} to your resume immediately. These are checked by every recruiter."
    elif absent:
        return f"Consider adding: {', '.join(absent)} to strengthen your online presence."
    else:
        return "✅ Excellent online presence! All major profile links are present."
